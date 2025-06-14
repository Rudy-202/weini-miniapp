from flask import Flask, request, jsonify, send_from_directory, redirect, url_for, send_file
from flask_cors import CORS
from models import db, User, Station, Task, InviteCode, Participant, Submission, GlobalSettings, VerificationCode, Feedback
from dotenv import load_dotenv
import os
import bcrypt
import jwt
from datetime import datetime, timedelta, timezone
from functools import wraps
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.utils import secure_filename
import uuid
from sqlalchemy import func, distinct
import io
import random
import string
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# 加载环境变量
load_dotenv()

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

# 创建Flask应用
app = Flask(__name__)

# 配置CORS
CORS(app, 
    origins=["http://localhost:5173", 
            "http://localhost:5174", 
            "http://192.168.1.167:5173", 
            "http://192.168.1.167:5174",
            "http://127.0.0.1:5173", 
            "http://127.0.0.1:5174",
            "http://localhost:5175",
            "http://localhost:3000"
            ],
    methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-Requested-With", "Accept"],
    supports_credentials=True,
    expose_headers=["Content-Type", "Authorization"],
    max_age=86400
)

# 配置请求限制器
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["2000 per day", "1000 per hour", "100 per minute"]
)

# 数据库配置
db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'weini.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# 初始化数据库
db.init_app(app)

# 从环境变量获取配置
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')  # 在生产环境中必须设置
app.config['ADMIN_EMAIL'] = os.getenv('ADMIN_EMAIL', 'admin@weini.com')
app.config['ADMIN_PASSWORD_HASH'] = os.getenv('ADMIN_PASSWORD_HASH')  # 存储加密后的密码

# 上传文件配置
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制上传文件大小为16MB
app.config['MAX_SUBMISSIONS'] = 5  # 每个任务最多5张图片

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 邮件发送配置
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.example.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME', 'email@example.com')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD', 'password')
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'True').lower() in ('true', '1', 't')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER', 'noreply@weini.com')

# 邮件发送函数
def send_email(to, subject, body):
    """发送电子邮件"""
    msg = MIMEMultipart()
    msg['From'] = app.config['MAIL_DEFAULT_SENDER']
    msg['To'] = to
    msg['Subject'] = subject
    
    msg.attach(MIMEText(body, 'html'))
    
    try:
        server = smtplib.SMTP(app.config['MAIL_SERVER'], app.config['MAIL_PORT'])
        if app.config['MAIL_USE_TLS']:
            server.starttls()
        server.login(app.config['MAIL_USERNAME'], app.config['MAIL_PASSWORD'])
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"邮件发送失败: {e}")
        return False

# 生成验证码
def generate_verification_code(length=6):
    """生成随机数字验证码"""
    return ''.join(random.choices(string.digits, k=length))

# 发送验证码API
@app.route('/api/admin/send-verification-code', methods=['POST'])
@limiter.limit("5 per hour")
def send_verification_code():
    """发送邮箱验证码"""
    data = request.get_json()
    if not data or not data.get('email'):
        return jsonify({"error": "请提供邮箱地址"}), 400
    
    email = data.get('email').strip()
    
    # 检查邮箱是否已注册
    existing_user = User.query.filter_by(email=email).first()
    if existing_user:
        return jsonify({"error": "该邮箱已注册"}), 409
    
    # 生成验证码
    code = generate_verification_code()
    
    # 设置过期时间为10分钟
    expires_at = datetime.utcnow() + timedelta(minutes=10)
    
    # 删除该邮箱的旧验证码
    VerificationCode.query.filter_by(email=email).delete()
    db.session.commit()
    
    # 保存新验证码
    new_code = VerificationCode(
        email=email,
        code=code,
        expires_at=expires_at
    )
    db.session.add(new_code)
    db.session.commit()
    
    # 发送邮件
    subject = "【微倪】站子管理员注册验证码"
    body = f"""
    <html>
    <body>
        <h2>微倪管理系统验证码</h2>
        <p>您好，</p>
        <p>感谢您注册微倪站子管理员账号。您的验证码是：</p>
        <h1 style="color: #4a6ee0;">{code}</h1>
        <p>此验证码将在10分钟内有效。</p>
        <p>如果您没有进行注册操作，请忽略此邮件。</p>
        <p>此致，</p>
        <p>微倪团队</p>
    </body>
    </html>
    """
    
    if send_email(email, subject, body):
        return jsonify({"message": "验证码已发送至您的邮箱"}), 200
    else:
        return jsonify({"error": "验证码发送失败，请稍后重试"}), 500

# 验证验证码API
@app.route('/api/admin/verify-code', methods=['POST'])
def verify_code():
    """验证邮箱验证码"""
    data = request.get_json()
    if not data or not data.get('email') or not data.get('code'):
        return jsonify({"error": "请提供邮箱和验证码"}), 400
    
    email = data.get('email').strip()
    code = data.get('code').strip()
    
    # 查找最新的有效验证码
    verification = VerificationCode.query.filter_by(
        email=email,
        verified=False
    ).order_by(VerificationCode.created_at.desc()).first()
    
    if not verification:
        return jsonify({"error": "未找到验证码，请重新获取"}), 404
    
    if verification.is_expired():
        return jsonify({"error": "验证码已过期，请重新获取"}), 401
    
    if verification.code != code:
        return jsonify({"error": "验证码不正确"}), 401
    
    # 验证成功，标记为已验证
    verification.verified = True
    db.session.commit()
    
    return jsonify({"message": "验证码验证成功"}), 200

# 管理员验证装饰器
def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        auth_header = request.headers.get('Authorization')
        
        if auth_header and auth_header.startswith('Bearer '):
            try:
                token = auth_header.split(" ")[1]
            except IndexError:
                return jsonify({'error': '无效的token格式'}), 401
        
        if not token:
            return jsonify({'error': '缺少token'}), 401
            
        try:
            data = jwt.decode(
                token, 
                app.config['SECRET_KEY'], 
                algorithms=["HS256"], 
                leeway=timedelta(seconds=10)
            )
            
            # 验证token中的用户角色为平台管理员或站子管理员
            if data.get('role') not in ['platform_admin', 'station_admin']:
                return jsonify({'error': '无效的管理员token (角色不匹配)'}), 401
            
            # 验证token是否过期
            expiration_time_timestamp = data.get('exp')
            if not expiration_time_timestamp:
                return jsonify({'error': '无效的token (缺少exp)'}), 401
            
            expiration_time = datetime.fromtimestamp(expiration_time_timestamp)
            current_time = datetime.utcnow()
            if expiration_time < current_time:
                return jsonify({'error': 'token已过期'}), 401
                    
        except jwt.ExpiredSignatureError:
            return jsonify({'error': 'token已过期'}), 401
        except jwt.InvalidTokenError as e:
            return jsonify({'error': f'无效的token: {str(e)}'}), 401
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Token验证服务器内部错误: {str(e)}'}), 500
            
        return f(*args, **kwargs)
    return decorated

# 站子管理员验证装饰器
def station_admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        auth_header = request.headers.get('Authorization')
        
        if auth_header and auth_header.startswith('Bearer '):
            try:
                token = auth_header.split(" ")[1]
            except IndexError:
                return jsonify({'error': '无效的token格式'}), 401
        
        if not token:
            return jsonify({'error': '缺少token'}), 401
            
        try:
            data = jwt.decode(
                token, 
                app.config['SECRET_KEY'], 
                algorithms=["HS256"], 
                leeway=timedelta(seconds=10)
            )
            
            # 验证token中的用户角色为站子管理员
            if data.get('role') != 'station_admin':
                return jsonify({'error': '无效的站子管理员token (角色不匹配)'}), 401
            
            # 验证token是否过期
            expiration_time_timestamp = data.get('exp')
            if not expiration_time_timestamp:
                return jsonify({'error': '无效的token (缺少exp)'}), 401
            
            expiration_time = datetime.fromtimestamp(expiration_time_timestamp)
            current_time = datetime.utcnow()
            if expiration_time < current_time:
                return jsonify({'error': 'token已过期'}), 401
            
            # 为视图函数设置当前用户ID
            kwargs['current_user_id'] = data.get('user_id')
                    
        except jwt.ExpiredSignatureError:
            return jsonify({'error': 'token已过期'}), 401
        except jwt.InvalidTokenError as e:
            return jsonify({'error': f'无效的token: {str(e)}'}), 401
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Token验证服务器内部错误: {str(e)}'}), 500
            
        return f(*args, **kwargs)
    return decorated

# 站子端注册API
@app.route('/api/admin/register', methods=['POST', 'OPTIONS'])
@limiter.limit("5 per hour")
def admin_register():
    # 处理OPTIONS请求
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    
    data = request.get_json()
    if not data or not data.get('email') or not data.get('password') or not data.get('code'):
        return jsonify({"error": "请提供邮箱、密码和验证码"}), 400

    email = data.get('email').strip()
    password = data.get('password')
    code = data.get('code').strip()
    username = data.get('username', email.split('@')[0])  # 默认使用邮箱前缀作为用户名

    # 检查邮箱是否已注册
    existing_user = User.query.filter_by(email=email).first()
    if existing_user:
        return jsonify({"error": "该邮箱已注册"}), 409
    
    # 验证验证码
    verification = VerificationCode.query.filter_by(
        email=email,
        verified=True
    ).order_by(VerificationCode.created_at.desc()).first()
    
    if not verification:
        return jsonify({"error": "请先验证邮箱"}), 403
    
    # 对密码进行哈希
    password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    # 创建用户
    new_user = User(
        email=email,
        password_hash=password_hash,
        role='station_admin',
        username=username,
        status='active',
        created_at=datetime.utcnow()
    )
    
    # 保存用户
    db.session.add(new_user)
    db.session.commit()
    
    # 创建默认站点
    new_station = Station(
        name=f"{username}的站点",
        owner_id=new_user.id,
        description="注册时自动创建的站点",
        status='active',
        created_at=datetime.utcnow()
    )
    
    db.session.add(new_station)
    db.session.commit()
    
    return jsonify({
        "message": "站子账号注册成功",
        "user_id": new_user.id,
        "station_id": new_station.id
    }), 201

# 站子端登录API
@app.route('/api/admin/login', methods=['POST'])
@limiter.limit("15 per minute")
def admin_login():
    data = request.get_json()
    if not data or not data.get('email') or not data.get('password'):
        return jsonify({"error": "请提供邮箱和密码"}), 400

    email = data.get('email')
    password = data.get('password')

    # 查询用户
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({'error': '无效的邮箱或密码'}), 401
    
    # 检查用户状态
    if user.status != 'active':
        return jsonify({'error': '账号已被禁用'}), 403
        
    # 验证密码
    if not bcrypt.checkpw(password.encode('utf-8'), user.password_hash.encode('utf-8')):
        return jsonify({'error': '无效的邮箱或密码'}), 401

    try:
        # 生成token
        current_utc_time = datetime.now(timezone.utc)
        expiration_time = current_utc_time + timedelta(hours=24)

        token_payload = {
            'user_id': user.id,
            'email': user.email,
            'role': user.role,
            'exp': expiration_time.timestamp(),
            'iat': current_utc_time.timestamp()
        }
        
        token = jwt.encode(
            token_payload,
            app.config['SECRET_KEY'],
            algorithm="HS256"
        )
        
        # 更新最后登录时间
        user.last_login = datetime.utcnow()
        db.session.commit()
        
        # 获取用户拥有的站点
        user_stations = Station.query.filter_by(owner_id=user.id).all()
        stations_data = [station.to_dict() for station in user_stations]
        
        response_data = {
            'token': token,
            'message': '登录成功',
            'user': {
                'id': user.id,
                'email': user.email,
                'username': user.username,
                'role': user.role
            },
            'stations': stations_data,
            'expires_in': 24 * 60 * 60  # 过期时间（秒）
        }
        
        return jsonify(response_data), 200

    except Exception as e:
        print(f"登录失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"登录失败: {str(e)}"}), 500

# 获取当前登录用户信息
@app.route('/api/admin/me', methods=['GET'])
@admin_required
def get_current_user():
    token = request.headers.get('Authorization').split(" ")[1]
    data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=["HS256"])
    
    user_id = data.get('user_id')
    user = User.query.get(user_id)
    
    if not user:
        return jsonify({'error': '用户不存在'}), 404
    
    # 获取用户拥有的站点
    user_stations = Station.query.filter_by(owner_id=user.id).all()
    stations_data = [station.to_dict() for station in user_stations]
    
    return jsonify({
        'user': user.to_dict(),
        'stations': stations_data
    }), 200

# 首页或健康检查
@app.route('/')
def index():
    return jsonify({
        'status': 'ok',
        'message': 'API服务正常运行',
        'timestamp': datetime.utcnow().isoformat()
    })

# 站子端创建任务
@app.route('/api/station/tasks', methods=['POST'])
@station_admin_required
def station_create_new_task(current_user_id):
    data = request.get_json()
    
    print(f"创建任务 - 用户ID: {current_user_id}, 请求数据: {data}")
    
    if not data:
        return jsonify({"error": "请求无效，缺少数据"}), 400
    
    # 验证必填字段
    required_fields = ['title', 'description', 'points', 'due_date', 'station_id', 'invite_code_id']
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        print(f"创建任务失败 - 缺少字段: {missing_fields}")
        return jsonify({'error': f'缺少必要字段: {", ".join(missing_fields)}'}), 400
    
    # 验证站点存在并且属于当前用户
    station = Station.query.filter_by(id=data['station_id'], owner_id=current_user_id).first()
    if not station:
        print(f"创建任务失败 - 站点不存在或无权限: 站点ID={data['station_id']}, 用户ID={current_user_id}")
        return jsonify({'error': '站点不存在或您没有权限操作此站点'}), 403
    
    # 验证邀请码
    invite_code_id = data.get('invite_code_id')
    invite_code = InviteCode.query.filter_by(
        id=invite_code_id, 
        station_id=station.id,
        status='active'
    ).first()
    
    if not invite_code:
        print(f"创建任务失败 - 邀请码不存在或不属于该站点: 邀请码ID={invite_code_id}, 站点ID={station.id}")
        return jsonify({'error': '邀请码不存在或不属于该站点'}), 400
    
    if invite_code.status != 'active':
        print(f"创建任务失败 - 邀请码已禁用: 邀请码ID={invite_code_id}, 状态={invite_code.status}")
        return jsonify({'error': '邀请码已禁用，无法创建任务'}), 400
    
    # 解析due_date字符串为datetime对象
    try:
        if data['due_date']:
            due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00'))
        else:
            due_date = None
    except ValueError as e:
        print(f"创建任务失败 - 截止日期格式无效: {data['due_date']}, 错误: {str(e)}")
        return jsonify({'error': '截止日期格式无效'}), 400
    
    # 创建新任务
    new_task = Task(
        station_id=station.id,
        invite_code_id=invite_code_id,
        title=data['title'],
        description=data['description'],
        points=data['points'],
        due_date=due_date,
        time_limit_mode=bool(due_date),
        flame_mode_enabled=data.get('flame_mode_enabled', True),
        bonus_points=data.get('bonus_points', 0),
        time_config=data.get('time_config', {
            'limit_hours': None,
            'bonus_points_on_time': 0,
            'fixed_start_time': None,
            'fixed_end_time': None
        }),
        is_focus_task=data.get('is_focus_task', False)
    )
    
    # 如果是焦点任务，取消该站点下其他焦点任务
    if new_task.is_focus_task:
        other_focus_tasks = Task.query.filter(
            Task.station_id == station.id,
            Task.is_focus_task == True,
            Task.status == 'active'
        ).all()
        
        for task in other_focus_tasks:
            task.is_focus_task = False
    
    # 保存任务
    try:
        db.session.add(new_task)
        db.session.commit()
        print(f"任务创建成功 - ID: {new_task.id}, 标题: {new_task.title}, 站点ID: {new_task.station_id}")
        return jsonify(new_task.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        print(f"任务创建失败 - 数据库错误: {str(e)}")
        return jsonify({'error': f'数据库错误: {str(e)}'}), 500

# 站子端获取任务列表
@app.route('/api/station/tasks', methods=['GET'])
@station_admin_required
def station_get_tasks(current_user_id):
    # 获取查询参数
    station_id = request.args.get('station_id')
    status = request.args.get('status', 'active')  # 默认获取活跃任务
    
    print(f"获取任务列表 - 用户ID: {current_user_id}, 站点ID: {station_id}, 状态: {status}")
    
    # 基础查询 - 获取当前用户拥有的站点的任务
    query = Task.query.join(Station).filter(Station.owner_id == current_user_id)
    
    # 如果指定了站点ID，则过滤特定站点的任务
    if station_id:
        query = query.filter(Task.station_id == station_id)
    
    # 根据状态过滤
    if status != 'all':
        query = query.filter(Task.status == status)
    
    # 按创建时间倒序排列
    tasks = query.order_by(Task.created_at.desc()).all()
    
    # 处理任务数据，添加参与人数统计
    result = []
    for task in tasks:
        task_dict = task.to_dict()
        
        # 查询参与人数
        participant_count = Participant.query.filter_by(task_id=task.id).count()
        task_dict['participant_count'] = participant_count
        
        result.append(task_dict)
    
    print(f"查询到 {len(result)} 个任务")
    
    # 打印每个任务的基本信息
    for idx, task in enumerate(tasks):
        print(f"任务 {idx+1}: ID={task.id}, 标题={task.title}, 站点ID={task.station_id}, 状态={task.status}")
    
    return jsonify(result), 200

# 站子端获取任务详情
@app.route('/api/station/tasks/<string:task_id>', methods=['GET'])
@station_admin_required
def station_get_task_detail(current_user_id, task_id):
    # 查询任务并验证权限
    task = Task.query.join(Station).filter(
        Task.id == task_id,
        Station.owner_id == current_user_id
    ).first()
    
    if not task:
        return jsonify({"error": "任务不存在或您没有权限查看"}), 404
    
    # 获取参与人数
    participant_count = Participant.query.filter_by(task_id=task.id).count()
    task_dict = task.to_dict()
    task_dict['participant_count'] = participant_count
    
    return jsonify(task_dict), 200

# 站子端更新任务
@app.route('/api/station/tasks/<string:task_id>', methods=['PUT'])
@station_admin_required
def station_update_task(current_user_id, task_id):
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "请求无效，缺少数据"}), 400
    
    # 查询任务并验证权限
    task = Task.query.join(Station).filter(
        Task.id == task_id,
        Station.owner_id == current_user_id
    ).first()
    
    if not task:
        return jsonify({"error": "任务不存在或您没有权限修改"}), 404
    
    # 更新任务字段
    if 'title' in data:
        task.title = data['title']
    
    if 'description' in data:
        task.description = data['description']
    
    if 'invite_code_id' in data:
        # 验证邀请码
        invite_code = InviteCode.query.filter_by(
            id=data['invite_code_id'], 
            station_id=task.station_id,
            status='active'
        ).first()
        
        if not invite_code:
            return jsonify({'error': '邀请码不存在或不属于该站点'}), 400
        
        if invite_code.status != 'active':
            return jsonify({'error': '邀请码已禁用，无法使用'}), 400
            
        task.invite_code_id = data['invite_code_id']
    
    if 'points' in data:
        task.points = data['points']
    
    if 'due_date' in data:
        try:
            if data['due_date']:
                task.due_date = datetime.fromisoformat(data['due_date'].replace('Z', '+00:00'))
                task.time_limit_mode = True
            else:
                task.due_date = None
                task.time_limit_mode = False
        except ValueError:
            return jsonify({'error': '截止日期格式无效'}), 400
    
    if 'time_limit_mode' in data:
        task.time_limit_mode = bool(data['time_limit_mode'])
    
    if 'flame_mode_enabled' in data:
        task.flame_mode_enabled = bool(data['flame_mode_enabled'])
    
    if 'bonus_points' in data:
        task.bonus_points = data['bonus_points']
    
    if 'time_config' in data:
        task.time_config = data['time_config']
    
    if 'status' in data:
        task.status = data['status']
        # 如果任务状态变为已完成，记录完成时间
        if data['status'] == 'completed' and not task.completed_at:
            task.completed_at = datetime.utcnow()
    
    if 'is_focus_task' in data:
        new_focus_status = bool(data['is_focus_task'])
        
        # 如果要关闭焦点任务，检查冷却期限制
        if task.is_focus_task and not new_focus_status:
            # 查询邀请码
            invite_code = InviteCode.query.get(task.invite_code_id)
            if invite_code and invite_code.last_focus_change:
                cooldown_until = invite_code.last_focus_change + timedelta(hours=24)
                now = datetime.utcnow()
                
                if now < cooldown_until:
                    remaining_time = cooldown_until - now
                    hours = int(remaining_time.total_seconds() // 3600)
                    minutes = int((remaining_time.total_seconds() % 3600) // 60)
                    return jsonify({
                        'error': f'焦点任务状态在24小时内不能取消。还需等待 {hours} 小时 {minutes} 分钟。'
                    }), 403
        
        # 如果设置为焦点任务，需要处理其他焦点任务
        if new_focus_status and not task.is_focus_task:
            # 获取站点最后一次焦点任务变更时间
            station = Station.query.get(task.station_id)
            last_change = InviteCode.query.filter_by(station_id=task.station_id, is_focus_enabled=True).first()
            
            if last_change and last_change.last_focus_change:
                cooldown_until = last_change.last_focus_change + timedelta(hours=24)
                now = datetime.utcnow()
                
                if now < cooldown_until:
                    return jsonify({'error': f'焦点任务状态在24小时内只能修改一次。上次修改时间：{last_change.last_focus_change.strftime("%Y-%m-%d %H:%M:%S UTC")}'}), 403
            
            # 取消其他焦点任务
            other_focus_tasks = Task.query.filter(
                Task.station_id == task.station_id,
                Task.is_focus_task == True,
                Task.id != task.id,
                Task.status == 'active'
            ).all()
            
            for other_task in other_focus_tasks:
                other_task.is_focus_task = False
            
            # 更新站点焦点任务变更时间
            if last_change:
                last_change.last_focus_change = datetime.utcnow()
        
        task.is_focus_task = new_focus_status
    
    # 保存更改
    db.session.commit()
    
    return jsonify(task.to_dict()), 200

# 站子端创建邀请码
@app.route('/api/station/invite-codes', methods=['POST'])
@station_admin_required
def station_create_invite_code(current_user_id):
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "请求无效，缺少数据"}), 400
    
    # 验证必填字段
    required_fields = ['station_id', 'description']
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        return jsonify({'error': f'缺少必要字段: {", ".join(missing_fields)}'}), 400
    
    # 验证站点存在并且属于当前用户
    station = Station.query.filter_by(id=data['station_id'], owner_id=current_user_id).first()
    if not station:
        return jsonify({'error': '站点不存在或您没有权限操作此站点'}), 403
    
    # 检查是否提供了自定义邀请码
    import re
    if 'custom_code' in data and data['custom_code']:
        custom_code = data['custom_code'].strip()
        
        # 验证邀请码格式：6-8位字母数字
        if not re.match(r'^[A-Za-z0-9]{6,8}$', custom_code):
            return jsonify({'error': '邀请码格式无效，必须是6-8位字母和数字的组合'}), 400
        
        # 检查邀请码是否已存在
        existing_code = InviteCode.query.filter_by(code=custom_code).first()
        if existing_code:
            return jsonify({'error': '该邀请码已被使用，请选择其他邀请码'}), 409
        
        code = custom_code
    else:
        # 生成随机邀请码
        import random
        import string
        
        # 生成8位随机邀请码
        while True:
            code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
            # 检查是否已存在
            existing_code = InviteCode.query.filter_by(code=code).first()
            if not existing_code:
                break
    
    # 创建新邀请码
    new_invite_code = InviteCode(
        code=code,
        station_id=station.id,
        description=data['description'],
        status='active',
        usage_limit=data.get('usage_limit', 'unlimited'),
        is_focus_enabled=data.get('is_focus_enabled', True)
    )
    
    # 保存邀请码
    db.session.add(new_invite_code)
    db.session.commit()
    
    return jsonify(new_invite_code.to_dict()), 201

# 站子端获取邀请码列表
@app.route('/api/station/invite-codes', methods=['GET'])
@station_admin_required
def station_get_invite_codes(current_user_id):
    # 获取查询参数
    station_id = request.args.get('station_id')
    
    # 基础查询 - 获取当前用户拥有的站点的邀请码
    query = InviteCode.query.join(Station).filter(Station.owner_id == current_user_id)
    
    # 如果指定了站点ID，则过滤特定站点的邀请码
    if station_id:
        query = query.filter(InviteCode.station_id == station_id)
    
    # 执行查询
    invite_codes = query.all()
    
    return jsonify([invite_code.to_dict() for invite_code in invite_codes]), 200

# 站子端获取邀请码详情
@app.route('/api/station/invite-codes/<string:invite_code_id>', methods=['GET'])
@station_admin_required
def station_get_invite_code_detail(current_user_id, invite_code_id):
    # 查询邀请码并验证权限
    invite_code = InviteCode.query.join(Station).filter(
        InviteCode.id == invite_code_id,
        Station.owner_id == current_user_id
    ).first()
    
    if not invite_code:
        return jsonify({"error": "邀请码不存在或您没有权限查看"}), 404
    
    return jsonify(invite_code.to_dict()), 200

# 站子端更新邀请码
@app.route('/api/station/invite-codes/<string:invite_code_id>', methods=['PUT'])
@station_admin_required
def station_update_invite_code(current_user_id, invite_code_id):
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "请求无效，缺少数据"}), 400
    
    # 查询邀请码并验证权限
    invite_code = InviteCode.query.join(Station).filter(
        InviteCode.id == invite_code_id,
        Station.owner_id == current_user_id
    ).first()
    
    if not invite_code:
        return jsonify({"error": "邀请码不存在或您没有权限修改"}), 404
    
    # 更新邀请码字段
    if 'description' in data:
        invite_code.description = data['description']
    
    if 'status' in data:
        invite_code.status = data['status']
    
    if 'usage_limit' in data:
        invite_code.usage_limit = data['usage_limit']
    
    if 'is_focus_enabled' in data:
        invite_code.is_focus_enabled = data['is_focus_enabled']
    
    # 保存更改
    db.session.commit()
    
    return jsonify(invite_code.to_dict()), 200

# 站子端获取站点列表
@app.route('/api/station/stations', methods=['GET'])
@station_admin_required
def station_get_stations(current_user_id):
    # 获取当前用户拥有的站点
    stations = Station.query.filter_by(owner_id=current_user_id).all()
    return jsonify([station.to_dict() for station in stations]), 200

# 站子端获取站点详情
@app.route('/api/station/stations/<string:station_id>', methods=['GET'])
@station_admin_required
def station_get_station_detail(current_user_id, station_id):
    # 查询站点并验证权限
    station = Station.query.filter_by(id=station_id, owner_id=current_user_id).first()
    
    if not station:
        return jsonify({"error": "站点不存在或您没有权限查看"}), 404
    
    return jsonify(station.to_dict()), 200

# 站子端更新站点信息
@app.route('/api/station/stations/<string:station_id>', methods=['PUT'])
@station_admin_required
def station_update_station(current_user_id, station_id):
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "请求无效，缺少数据"}), 400
    
    # 查询站点并验证权限
    station = Station.query.filter_by(id=station_id, owner_id=current_user_id).first()
    
    if not station:
        return jsonify({"error": "站点不存在或您没有权限修改"}), 404
    
    # 更新站点字段
    if 'name' in data:
        station.name = data['name']
    
    if 'description' in data:
        station.description = data['description']
    
    if 'logo_url' in data:
        station.logo_url = data['logo_url']
    
    if 'settings' in data:
        station.settings = data['settings']
    
    # 保存更改
    station.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify(station.to_dict()), 200

# 站子端创建新站点
@app.route('/api/station/stations', methods=['POST'])
@station_admin_required
def station_create_station(current_user_id):
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "请求无效，缺少数据"}), 400
    
    # 验证必填字段
    required_fields = ['name']
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        return jsonify({'error': f'缺少必要字段: {", ".join(missing_fields)}'}), 400
    
    # 创建新站点
    new_station = Station(
        name=data['name'],
        owner_id=current_user_id,
        description=data.get('description', ''),
        logo_url=data.get('logo_url', ''),
        settings=data.get('settings', {}),
        status='active'
    )
    
    # 保存站点
    db.session.add(new_station)
    db.session.commit()
    
    return jsonify(new_station.to_dict()), 201

# 检查邀请码的焦点任务状态API
@app.route('/api/station/invite-codes/<string:invite_code>/focus-status', methods=['GET'])
def get_invite_code_focus_status(invite_code):
    try:
        # 查找邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            return jsonify({"error": "邀请码不存在或未激活"}), 404
        
        # 检查该邀请码是否有焦点任务
        focus_tasks = Task.query.filter_by(
            invite_code_id=invite_code_obj.id, 
            is_focus_task=True,
            status='active'
        ).all()
        
        has_focus_task = len(focus_tasks) > 0
        focus_task_info = focus_tasks[0] if has_focus_task else None
        
        # 检查冷却状态
        now_utc = datetime.utcnow()
        last_change = invite_code_obj.last_focus_change
        
        # 计算冷却截止时间点
        cooldown_until_dt = last_change + timedelta(hours=24) if last_change else datetime(1970, 1, 1)
        
        # 判断是否在冷却期
        is_in_cooldown = last_change and now_utc < cooldown_until_dt
        
        # 计算剩余冷却时间
        cooldown_remaining_seconds = 0
        if is_in_cooldown:
            cooldown_remaining_seconds = int((cooldown_until_dt - now_utc).total_seconds())
        
        # 构建响应
        response = {
            "invite_code": invite_code,
            "has_focus_task": has_focus_task,
            "is_in_cooldown": is_in_cooldown,
            "cooldown_remaining_seconds": cooldown_remaining_seconds,
            "last_change_time": last_change.isoformat() if last_change else None,
            "cooldown_until_time": cooldown_until_dt.isoformat() if is_in_cooldown else None,
            "current_time": now_utc.isoformat()
        }
        
        # 如果有焦点任务，添加基本信息
        if has_focus_task:
            response["focus_task"] = {
                "id": focus_task_info.id,
                "title": focus_task_info.title,
                "created_at": focus_task_info.created_at.isoformat()
            }
        
        return jsonify(response)
    except Exception as e:
        print(f"获取焦点任务状态时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"获取焦点任务状态失败: {str(e)}"}), 500

# 粉丝根据邀请码查询任务
@app.route('/api/fan/tasks', methods=['GET'])
def get_fan_tasks():
    invite_code = request.args.get('invite_code')
    
    print(f"获取粉丝任务列表 - 邀请码: {invite_code}")
    
    if not invite_code:
        return jsonify({"error": "请提供邀请码"}), 400
    
    try:
        # 查找邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            print(f"邀请码不存在或已失效: {invite_code}")
            return jsonify({"error": "邀请码不存在或已失效"}), 404
        
        print(f"邀请码已找到 - ID: {invite_code_obj.id}, 站点ID: {invite_code_obj.station_id}")
        
        # 获取该邀请码下所有活跃任务
        tasks = Task.query.filter_by(
            invite_code_id=invite_code_obj.id,
            status='active'
        ).all()
        
        print(f"查询到 {len(tasks)} 个活跃任务")
        
        # 转换为JSON
        tasks_data = []
        for task in tasks:
            try:
                # 查询任务的总提交次数
                submission_count = db.session.query(func.sum(Participant.submission_count)).filter(
                    Participant.task_id == task.id
                ).scalar() or 0
                
                task_dict = {
                    'id': task.id,
                    'title': task.title,
                    'description': task.description,
                    'points': task.points,
                    'due_date': task.due_date.isoformat() if task.due_date else None,
                    'is_focus_task': task.is_focus_task,
                    'flame_mode_enabled': task.flame_mode_enabled,
                    'created_at': task.created_at.isoformat() if task.created_at else None,
                    'bonus_points': task.bonus_points,
                    'submission_count': submission_count  # 添加提交次数
                }
                tasks_data.append(task_dict)
            except Exception as e:
                print(f"处理任务数据时出错 - 任务ID: {task.id}, 错误: {str(e)}")
        
        print(f"成功处理 {len(tasks_data)} 个任务数据")
        return jsonify(tasks_data), 200
    
    except Exception as e:
        print(f"获取任务列表失败 - 邀请码: {invite_code}, 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"获取任务失败: {str(e)}"}), 500

# 添加单个任务详情的API路由
@app.route('/api/fan/tasks/<string:task_id>', methods=['GET'])
def get_fan_task_detail(task_id):
    """获取单个任务详情（粉丝视角）"""
    invite_code = request.args.get('invite_code')
    nickname = request.args.get('nickname', '')
    
    print(f"获取任务详情 - 任务ID: {task_id}, 邀请码: {invite_code}, 昵称: {nickname}")
    
    if not invite_code:
        return jsonify({'error': '缺少必要参数：邀请码'}), 400
    
    try:
        # 验证邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            print(f"邀请码无效: {invite_code}")
            return jsonify({'error': '无效的邀请码'}), 403
        
        print(f"邀请码已验证 - ID: {invite_code_obj.id}, 站点ID: {invite_code_obj.station_id}")
        
        # 查询任务详情
        task = Task.query.get(task_id)
        if not task:
            print(f"任务不存在: {task_id}")
            return jsonify({'error': '任务不存在'}), 404
        
        print(f"任务已找到 - 标题: {task.title}, 站点ID: {task.station_id}")
        
        # 检查任务对应的站点是否匹配邀请码
        if task.station_id != invite_code_obj.station_id:
            print(f"任务站点不匹配 - 任务站点ID: {task.station_id}, 邀请码站点ID: {invite_code_obj.station_id}")
            return jsonify({'error': '无权访问此任务'}), 403
        
        # 查询用户参与记录
        participation = None
        if nickname:
            participation = Participant.query.filter_by(
                name=nickname,
                task_id=task_id
            ).first()
            if participation:
                print(f"找到参与记录 - 提交次数: {participation.submission_count}, 积分: {participation.points_earned}")
            else:
                print(f"未找到参与记录 - 昵称: {nickname}")
        
        # 构建响应
        response_data = {
            'id': task.id,
            'title': task.title,
            'description': task.description,
            'points': task.points,
            'bonus_points': task.bonus_points,
            'created_at': task.created_at.isoformat() if task.created_at else None,
            'due_date': task.due_date.isoformat() if task.due_date else None,
            'is_focus_task': task.is_focus_task,
            'display_focus_icon': '🌟' if task.is_focus_task else None,
            'status': task.status,
            'time_limit_mode': task.due_date is not None,
            'flame_mode_enabled': task.flame_mode_enabled,
            'has_participated': participation is not None,
            'submission_count': participation.submission_count if participation else 0
        }
        
        print(f"任务详情返回成功 - 任务ID: {task_id}")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"获取任务详情失败 - 任务ID: {task_id}, 邀请码: {invite_code}, 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'获取任务详情失败: {str(e)}'}), 500

# 任务提交API
@app.route('/api/fan/tasks/<string:task_id>/submit', methods=['POST'])
def submit_task_for_task_id(task_id):
    """提交指定ID的任务完成证明"""
    try:
        print("="*50)
        print(f"开始处理任务提交 - 任务ID: {task_id}")
        
        # 获取必要参数
        nickname = request.form.get('nickname')
        comment = request.form.get('comment', '')
        invite_code = request.form.get('invite_code')
        
        # 详细的请求信息
        print(f"请求方法: {request.method}")
        print(f"Content-Type: {request.headers.get('Content-Type')}")
        print(f"请求表单数据: {request.form}")
        print(f"请求文件: {list(request.files.keys())}")
        
        # 参数验证
        if not nickname or not nickname.strip():
            return jsonify({'error': '请提供昵称'}), 400
        
        if not invite_code or not invite_code.strip():
            return jsonify({'error': '请提供邀请码'}), 400
        
        # 验证邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            print(f"邀请码无效: {invite_code}")
            return jsonify({'error': '无效的邀请码'}), 403
        
        print(f"邀请码验证通过: {invite_code}")
        
        # 检查任务是否存在
        task = Task.query.get(task_id)
        if not task:
            print(f"任务不存在: {task_id}")
            return jsonify({'error': '任务不存在'}), 404
        
        print(f"找到任务: {task.title}")
        
        # 检查任务是否属于邀请码对应的站点
        if task.station_id != invite_code_obj.station_id:
            print(f"任务与邀请码不匹配 - 任务站点: {task.station_id}, 邀请码站点: {invite_code_obj.station_id}")
            return jsonify({'error': '无权访问此任务'}), 403
            
        # 检查任务是否已截止
        if task.due_date and task.due_date < datetime.utcnow() and not task.flame_mode_enabled:
            print(f"任务已截止 - 截止时间: {task.due_date.isoformat()}")
            return jsonify({'error': '任务已截止'}), 400
            
        # 检查是否有文件上传
        uploaded_images = []
        
        # 处理图片上传 - 尝试多种可能的字段名称
        possible_fields = ['images', 'image', 'files', 'file']
        
        for field in possible_fields:
            if field in request.files:
                field_files = request.files.getlist(field)
                if field_files and field_files[0].filename:
                    uploaded_images.extend(field_files)
                    print(f"从字段'{field}'获取到{len(field_files)}个文件")
        
        # 如果没有找到文件，尝试遍历所有文件字段
        if not uploaded_images:
            for key in request.files:
                files = request.files.getlist(key)
                if files and files[0].filename:
                    uploaded_images.extend(files)
                    print(f"从额外字段'{key}'获取到{len(files)}个文件")
        
        # 最终检查是否有图片
        if not uploaded_images:
            print("没有找到有效的图片文件")
            return jsonify({'error': '请至少上传一张图片'}), 400
        
        print(f"共获取到{len(uploaded_images)}张图片")
        
        # 处理图片保存
        image_paths = []
        max_file_size = 5 * 1024 * 1024  # 5MB最大文件大小
        
        for image in uploaded_images:
            try:
                print(f"处理图片: {image.filename}")
                
                # 检查文件名
                if not image.filename or not image.filename.strip():
                    print("文件名为空，跳过")
                    continue
                
                # 检查文件类型
                if not allowed_file(image.filename):
                    print(f"文件类型不支持: {image.filename}")
                    continue
                
                # 检查文件大小
                image.seek(0, os.SEEK_END)
                file_size = image.tell()
                image.seek(0)
                if file_size > max_file_size:
                    print(f"文件大小超过限制: {file_size} 字节")
                    continue
                
                # 生成安全的文件名
                filename = secure_filename(image.filename)
                # 添加时间戳和随机字符串，确保文件名唯一
                timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                unique_id = str(uuid.uuid4())[:8]
                new_filename = f"{timestamp}_{unique_id}_{filename}"
                
                # 确保目录存在
                upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'task_submissions')
                os.makedirs(upload_dir, exist_ok=True)
                
                # 保存文件
                filepath = os.path.join(upload_dir, new_filename)
                print(f"保存文件到: {filepath}")
                image.save(filepath)
                image_paths.append(filepath)
                print(f"图片保存成功: {filepath}")
            except Exception as e:
                print(f"处理图片时出错: {str(e)}")
                continue
                
        if not image_paths:
            return jsonify({'error': '未能成功保存任何图片，请检查图片格式和大小'}), 400
        
        print(f"成功保存了{len(image_paths)}张图片")
        
        # 查找或创建参与记录
        participant = Participant.query.filter_by(
            name=nickname,
            task_id=task_id
        ).first()
        
        if not participant:
            print(f"创建新的参与者记录: {nickname}")
            participant = Participant(
                name=nickname,
                task_id=task_id,
                submission_count=0,
                points_earned=0,
                total_points_for_task=0
            )
            db.session.add(participant)
            db.session.flush()  # Flush to get the participant.id before creating Submission
        else:
            print(f"找到现有参与者: {nickname}, 提交次数: {participant.submission_count}")
            
        # 创建提交记录
        submission = Submission(
            participant_id=participant.id,
            comment=comment,
            image_urls=image_paths,
            is_abnormal=False,  # 默认不是异常提交
            points_earned=0     # 初始分数为0
        )
        
        db.session.add(submission)
        print(f"创建提交记录: {submission.id}")
        
        # 更新参与计数
        participant.submission_count += 1
        
        # 如果是首次提交，增加积分
        points_earned = 0
        first_time = participant.submission_count == 1
        
        if first_time or task.flame_mode_enabled:
            # 基础积分
            points_earned = task.points
            
            # 如果是焦点任务，积分翻倍
            if task.is_focus_task:
                points_earned *= 2
                
            # 如果有限时奖励且在截止日期前提交
            if task.bonus_points and task.due_date and task.due_date > datetime.utcnow():
                if task.is_focus_task:
                    points_earned += (task.bonus_points * 2)
                else:
                    points_earned += task.bonus_points
            
            # 更新提交记录的积分
            submission.points_earned = points_earned
                    
            # 更新参与者积分
            participant.points_earned += points_earned
            participant.total_points_for_task += points_earned
            print(f"积分更新: +{points_earned}分, 总计: {participant.points_earned}分")
            
        db.session.commit()
        print(f"提交完成 - 任务ID: {task_id}, 参与者: {nickname}, 获得积分: {points_earned}")
        
        return jsonify({
            'success': True,
            'message': '任务提交成功',
            'points': points_earned,
            'submission_id': submission.id,
            'total_points': participant.points_earned
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"提交任务失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'提交任务失败: {str(e)}'}), 500

# 验证文件扩展名是否合法
def allowed_file(filename):
    """验证文件扩展名是否在允许的列表中"""
    if not filename:
        return False
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 安全的文件名处理
def secure_filename(filename):
    """生成安全的文件名，避免路径遍历攻击"""
    from werkzeug.utils import secure_filename as _secure_filename
    return _secure_filename(filename)

# 任务鼓励内容
@app.route('/api/fan/encouragement', methods=['GET'])
def get_fan_encouragement():
    """获取任务的鼓励内容。"""
    try:
        task_id = request.args.get('task_id')
        invite_code = request.args.get('invite_code')
        
        if not task_id or not invite_code:
            return jsonify({'error': '缺少必要参数'}), 400
        
        # 验证邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            return jsonify({'error': '无效的邀请码'}), 403
        
        # 获取任务信息
        task = Task.query.get(task_id)
        if not task:
            return jsonify({'error': '任务不存在'}), 404
            
        # 检查任务是否属于邀请码对应的站点
        if task.station_id != invite_code_obj.station_id:
            return jsonify({'error': '无权访问此任务'}), 403
        
        # 获取全局鼓励设置
        global_settings = GlobalSettings.query.first()
        
        # 构建响应
        response = {
            'title': '任务完成！',
            'message': '恭喜你成功完成任务！感谢你的付出和努力~',
            'image_url': None
        }
        
        # 如果有特定任务的鼓励内容，优先使用
        if task.encouragement_message:
            response['message'] = task.encouragement_message
            
        if task.encouragement_image_url:
            response['image_url'] = task.encouragement_image_url
            
        # 如果任务没有鼓励内容，使用全局设置
        elif global_settings:
            if not task.encouragement_message and global_settings.default_encouragement_message:
                response['message'] = global_settings.default_encouragement_message
                
            if not task.encouragement_image_url and global_settings.default_encouragement_image_url:
                response['image_url'] = global_settings.default_encouragement_image_url
                
        return jsonify(response)
        
    except Exception as e:
        print(f"获取鼓励内容失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'获取鼓励内容失败: {str(e)}'}), 500

# 获取全局鼓励设置
@app.route('/api/fan/global-encouragement-settings', methods=['GET'])
def get_fan_global_encouragement_settings():
    """获取全局鼓励设置"""
    try:
        invite_code = request.args.get('invite_code')
        
        if not invite_code:
            # 如果没有提供邀请码，尝试获取默认设置
            global_settings = GlobalSettings.query.first()
            
            if not global_settings:
                return jsonify({
                    'default_encouragement_message': '恭喜你成功完成任务！感谢你的付出和努力~',
                    'default_encouragement_image_url': None
                })
                
            # 构建响应
            response = {
                'default_encouragement_message': global_settings.default_encouragement_message or '恭喜你成功完成任务！感谢你的付出和努力~',
                'default_encouragement_image_url': global_settings.default_encouragement_image_url
            }
            
            return jsonify(response)
        
        # 验证邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            return jsonify({'error': '无效的邀请码'}), 403
            
        # 获取站点全局设置
        global_settings = GlobalSettings.query.first()
        
        if not global_settings:
            return jsonify({
                'default_encouragement_message': '恭喜你成功完成任务！感谢你的付出和努力~',
                'default_encouragement_image_url': None
            })
            
        # 构建响应
        response = {
            'default_encouragement_message': global_settings.default_encouragement_message or '恭喜你成功完成任务！感谢你的付出和努力~',
            'default_encouragement_image_url': global_settings.default_encouragement_image_url
        }
        
        return jsonify(response)
        
    except Exception as e:
        print(f"获取全局鼓励设置失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'获取全局鼓励设置失败: {str(e)}'}), 500

# 任务排行榜
@app.route('/api/fan/leaderboard/<string:task_id>', methods=['GET'])
def get_task_leaderboard(task_id):
    """获取指定任务的排行榜数据"""
    try:
        invite_code = request.args.get('invite_code')
        fan_nickname = request.args.get('fan_nickname', '')
        
        if not invite_code:
            return jsonify({'error': '缺少必要参数：邀请码'}), 400
            
        # 验证邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            return jsonify({'error': '无效的邀请码'}), 403
            
        # 查询任务
        task = Task.query.get(task_id)
        if not task:
            return jsonify({'error': '任务不存在'}), 404
            
        # 检查任务是否属于邀请码对应的站点
        if task.station_id != invite_code_obj.station_id:
            return jsonify({'error': '无权访问此任务排行榜'}), 403
        
        # 确保任何待处理的数据库事务都已提交    
        db.session.commit()
            
        # 查询任务参与者
        participants = Participant.query.filter_by(
            task_id=task_id
        ).order_by(Participant.points_earned.desc()).all()
        
        # 构建排行榜数据
        leaderboard_data = []
        self_info = None
        
        for index, participant in enumerate(participants):
            # 基本信息
            participant_data = {
                'nickname': participant.name,
                'points': participant.points_earned,
                'completed_tasks': participant.submission_count,
                'rank': index + 1,
                'has_focus_task_completed': task.is_focus_task and participant.submission_count > 0
            }
            
            leaderboard_data.append(participant_data)
            
            # 如果是当前用户，记录自己的信息
            if participant.name == fan_nickname:
                self_info = participant_data
                
        # 如果没找到自己的信息，但提供了昵称
        if not self_info and fan_nickname:
            self_info = {
                'nickname': fan_nickname,
                'points': 0,
                'completed_tasks': 0,
                'rank': len(participants) + 1,
                'has_focus_task_completed': False
            }
        
        # 如果请求中包含用户昵称参数，直接返回该用户的排名信息
        if fan_nickname:
            return jsonify({
                "leaderboard": leaderboard_data,
                "user_info": self_info
            })
            
        return jsonify(leaderboard_data)
        
    except Exception as e:
        print(f"获取任务排行榜失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'获取任务排行榜失败: {str(e)}'}), 500

# 通用排行榜
@app.route('/api/fan/leaderboard', methods=['GET'])
def get_fan_leaderboard():
    """获取排行榜数据（总榜、日榜、焦点榜）
    
    排行榜计算逻辑:
    - 总榜: 显示所有历史累计积分
    - 日榜: 仅显示当日提交任务获得的积分
    - 焦点榜: 仅统计焦点任务获得的积分
    """
    try:
        invite_code = request.args.get('invite_code')
        leaderboard_type = request.args.get('type', 'overall')
        task_id = request.args.get('task_id')
        
        if not invite_code:
            return jsonify({'error': '缺少必要参数：邀请码'}), 400
            
        # 验证邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code, status='active').first()
        if not invite_code_obj:
            return jsonify({'error': '无效的邀请码'}), 403
            
        # 获取站点ID
        station_id = invite_code_obj.station_id
        
        # 确保任何待处理的数据库事务都已提交
        db.session.commit()
        
        # 根据类型查询不同的排行榜
        if leaderboard_type == 'task' and task_id:
            # 如果是任务排行榜，重定向到任务特定的排行榜API
            from flask import redirect, url_for
            
            # 构建查询参数
            query_params = request.args.copy()
            # 确保任务ID不被添加为查询参数
            if 'task_id' in query_params:
                del query_params['task_id']
                
            # 重定向到特定的任务排行榜API
            return redirect(url_for('get_task_leaderboard', task_id=task_id, **query_params))
        
        elif leaderboard_type == 'daily':
            # 日榜 - 今日参与的任务积分
            today = datetime.now().date()
            tomorrow = today + timedelta(days=1)
            
            # 获取今日0点到明日0点的时间范围
            today_start = datetime.combine(today, datetime.min.time()).replace(tzinfo=timezone.utc)
            today_end = datetime.combine(tomorrow, datetime.min.time()).replace(tzinfo=timezone.utc)
            
            # 使用Submission表直接统计今日提交的积分
            # 这样可以确保只统计今日提交所获得的积分
            daily_participants = db.session.query(
                Participant.name,
                func.sum(Submission.points_earned).label('total_points'),
                func.count(distinct(Submission.id)).label('submission_count')
            ).join(
                Submission, Submission.participant_id == Participant.id
            ).join(
                Task, Task.id == Participant.task_id
            ).filter(
                Submission.submitted_at >= today_start,
                Submission.submitted_at < today_end,
                Task.station_id == station_id
            ).group_by(
                Participant.name
            ).order_by(
                func.sum(Submission.points_earned).desc()
            ).all()
            
            # 构建响应数据
            daily_leaderboard = []
            for idx, (name, points, submission_count) in enumerate(daily_participants):
                daily_leaderboard.append({
                    'nickname': name,
                    'points': points,
                    'completed_tasks': submission_count,
                    'rank': idx + 1
                })
                
            return jsonify(daily_leaderboard)
            
        elif leaderboard_type == 'focus':
            # 焦点榜 - 焦点任务积分
            focus_participants = db.session.query(
                Participant.name,
                func.sum(Participant.points_earned).label('total_points'),
                func.count(Participant.id).label('completed_task_count')
            ).join(
                Task, Task.id == Participant.task_id
            ).filter(
                Task.is_focus_task == True,
                Task.station_id == station_id
            ).group_by(
                Participant.name
            ).order_by(
                func.sum(Participant.points_earned).desc()
            ).all()
            
            # 构建响应数据
            focus_leaderboard = []
            for idx, (name, points, task_count) in enumerate(focus_participants):
                focus_leaderboard.append({
                    'nickname': name,
                    'points': points,
                    'completed_tasks': task_count,
                    'rank': idx + 1,
                    'has_focus_task_completed': True
                })
                
            return jsonify(focus_leaderboard)
            
        else:
            # 总榜 - 所有任务总积分
            overall_participants = db.session.query(
                Participant.name,
                func.sum(Participant.points_earned).label('total_points'),
                func.count(Participant.id).label('completed_task_count')
            ).join(
                Task, Task.id == Participant.task_id
            ).filter(
                Task.station_id == station_id
            ).group_by(
                Participant.name
            ).order_by(
                func.sum(Participant.points_earned).desc()
            ).all()
            
            # 查询已完成焦点任务的用户
            focus_task_completers = db.session.query(
                distinct(Participant.name)
            ).join(
                Task, Task.id == Participant.task_id
            ).filter(
                Task.is_focus_task == True,
                Task.station_id == station_id,
                Participant.submission_count > 0
            ).all()
            
            focus_completers = set([name[0] for name in focus_task_completers])
            
            # 构建响应数据
            overall_leaderboard = []
            for idx, (name, points, task_count) in enumerate(overall_participants):
                overall_leaderboard.append({
                    'nickname': name,
                    'points': points,
                    'completed_tasks': task_count,
                    'rank': idx + 1,
                    'has_focus_task_completed': name in focus_completers
                })
                
            return jsonify(overall_leaderboard)
                
    except Exception as e:
        print(f"获取排行榜失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'获取排行榜失败: {str(e)}'}), 500

# 测试用的API路由
@app.route('/api/fan/tasks/test-task-id/submit', methods=['POST'])
def test_upload():
    """测试图片上传功能的API端点"""
    print("="*50)
    print("测试图片上传接口被调用")
    
    try:
        # 获取表单数据
        nickname = request.form.get('nickname', '测试用户')
        comment = request.form.get('comment', '')
        invite_code = request.form.get('invite_code', 'TEST123')
        
        print(f"昵称: {nickname}")
        print(f"备注: {comment}")
        print(f"邀请码: {invite_code}")
        
        # 获取文件信息
        print(f"请求文件字段列表: {list(request.files.keys())}")
        
        # 检查'images'字段
        if 'images' in request.files:
            image_files = request.files.getlist('images')
            print(f"images文件数量: {len(image_files)}")
            
            for i, img in enumerate(image_files):
                print(f"图片{i+1}: {img.filename} - {img.content_type} - {img.content_length or '未知'}字节")
            
            # 保存文件
            saved_files = []
            for image in image_files:
                if image and image.filename:
                    filename = secure_filename(image.filename)
                    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                    unique_id = str(uuid.uuid4())[:8]
                    new_filename = f"{timestamp}_{unique_id}_{filename}"
                    
                    # 确保目录存在
                    test_upload_dir = os.path.join('uploads', 'test_uploads')
                    os.makedirs(test_upload_dir, exist_ok=True)
                    
                    # 保存文件
                    filepath = os.path.join(test_upload_dir, new_filename)
                    image.save(filepath)
                    saved_files.append({
                        "original_name": image.filename,
                        "saved_as": new_filename,
                        "path": filepath
                    })
                    print(f"保存文件: {filepath}")
            
            return jsonify({
                "success": True,
                "message": f"成功上传{len(saved_files)}个文件",
                "files": saved_files,
                "form_data": {
                    "nickname": nickname,
                    "comment": comment,
                    "invite_code": invite_code
                }
            })
        else:
            # 如果没有images字段，检查其他可能的字段
            all_files = []
            for field_name in request.files:
                field_files = request.files.getlist(field_name)
                print(f"字段{field_name}中有{len(field_files)}个文件")
                all_files.extend([(field_name, f.filename) for f in field_files])
            
            return jsonify({
                "success": False,
                "error": "未找到'images'字段",
                "found_fields": list(request.files.keys()),
                "all_files": all_files
            }), 400
    
    except Exception as e:
        import traceback
        error_msg = f"图片上传测试失败: {str(e)}"
        print(error_msg)
        traceback.print_exc()
        return jsonify({"success": False, "error": error_msg}), 500

@app.route('/api/station/tasks/<string:task_id>/submissions', methods=['GET'])
@station_admin_required
def station_get_task_submissions(current_user_id, task_id):
    """获取指定任务的所有提交记录"""
    # 获取任务
    task = Task.query.filter_by(id=task_id).first()
    if not task:
        return jsonify({'error': '任务不存在'}), 404
        
    # 验证权限(检查任务是否属于管理员的站点)
    admin_stations = Station.query.filter_by(owner_id=current_user_id).all()
    admin_station_ids = [station.id for station in admin_stations]
    
    if task.station_id not in admin_station_ids:
        return jsonify({'error': '无权访问此任务'}), 403
    
    # 获取任务的所有参与者
    participants = Participant.query.filter_by(task_id=task_id).all()
    
    # 收集所有提交
    submissions_data = []
    for participant in participants:
        submissions = Submission.query.filter_by(participant_id=participant.id).all()
        for submission in submissions:
            submissions_data.append({
                'id': submission.id,
                'participant_id': participant.id,
                'participant_name': participant.name,
                'submitted_at': submission.submitted_at.isoformat(),
                'points_earned': submission.points_earned,
                'is_abnormal': submission.is_abnormal,
                'abnormal_reason': submission.abnormal_reason if submission.is_abnormal else None,
                'image_preview': submission.image_urls[0] if submission.image_urls else None,
                'submission_count': len(submission.image_urls) if submission.image_urls else 0
            })
    
    # 按提交时间排序，最新的在前
    submissions_data.sort(key=lambda x: x['submitted_at'], reverse=True)
    
    return jsonify(submissions_data), 200

@app.route('/api/station/submissions/<string:submission_id>', methods=['GET'])
@station_admin_required
def station_get_submission_detail(current_user_id, submission_id):
    """获取单个提交的详细信息"""
    # 获取提交
    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        return jsonify({'error': '提交记录不存在'}), 404
    
    # 获取参与者
    participant = Participant.query.filter_by(id=submission.participant_id).first()
    if not participant:
        return jsonify({'error': '参与者记录不存在'}), 404
    
    # 获取任务
    task = Task.query.filter_by(id=participant.task_id).first()
    if not task:
        return jsonify({'error': '任务记录不存在'}), 404
    
    # 验证权限
    admin_stations = Station.query.filter_by(owner_id=current_user_id).all()
    admin_station_ids = [station.id for station in admin_stations]
    
    if task.station_id not in admin_station_ids:
        return jsonify({'error': '无权访问此提交记录'}), 403
    
    # 获取标记人信息（如果已标记）
    marker_info = None
    if submission.marked_by:
        marker = User.query.filter_by(id=submission.marked_by).first()
        if marker:
            marker_info = {
                'id': marker.id,
                'username': marker.username,
                'email': marker.email
            }
    
    # 构建详细信息
    submission_detail = {
        'id': submission.id,
        'participant': {
            'id': participant.id,
            'name': participant.name,
            'joined_at': participant.joined_at.isoformat() if participant.joined_at else None,
            'submission_count': participant.submission_count,
            'points_earned': participant.points_earned
        },
        'task': {
            'id': task.id,
            'title': task.title,
            'station_id': task.station_id
        },
        'submitted_at': submission.submitted_at.isoformat(),
        'points_earned': submission.points_earned,
        'comment': submission.comment,
        'is_abnormal': submission.is_abnormal,
        'abnormal_reason': submission.abnormal_reason,
        'marked_at': submission.marked_at.isoformat() if submission.marked_at else None,
        'marked_by': marker_info,
        'image_urls': submission.image_urls
    }
    
    return jsonify(submission_detail), 200

@app.route('/api/station/submissions/<string:submission_id>/mark-abnormal', methods=['POST'])
@station_admin_required
def station_mark_abnormal_submission(current_user_id, submission_id):
    """将提交标记为异常，并扣除相应积分"""
    data = request.get_json()
    if not data:
        return jsonify({'error': '请提供标记原因'}), 400
    
    reason = data.get('reason', '未提供原因')
    if not reason.strip():
        return jsonify({'error': '请提供标记原因'}), 400
    
    # 获取提交
    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        return jsonify({'error': '提交记录不存在'}), 404
    
    # 检查是否已标记
    if submission.is_abnormal:
        return jsonify({'error': '该提交已被标记为异常'}), 400
    
    # 获取参与者
    participant = Participant.query.filter_by(id=submission.participant_id).first()
    if not participant:
        return jsonify({'error': '参与者记录不存在'}), 404
    
    # 获取任务
    task = Task.query.filter_by(id=participant.task_id).first()
    if not task:
        return jsonify({'error': '任务记录不存在'}), 404
    
    # 验证权限
    admin_stations = Station.query.filter_by(owner_id=current_user_id).all()
    admin_station_ids = [station.id for station in admin_stations]
    
    if task.station_id not in admin_station_ids:
        return jsonify({'error': '无权操作此提交记录'}), 403
    
    try:
        # 开始数据库事务
        # 标记为异常
        submission.is_abnormal = True
        submission.abnormal_reason = reason
        submission.marked_by = current_user_id
        submission.marked_at = datetime.utcnow()
        
        # 记录原始积分用于返回
        points_to_deduct = submission.points_earned
        
        # 如果有积分则扣除
        if points_to_deduct > 0:
            # 从参与者总积分中扣除
            participant.points_earned = max(0, participant.points_earned - points_to_deduct)
            participant.total_points_for_task = max(0, participant.total_points_for_task - points_to_deduct)
            
            # 将提交的积分设为0
            submission.points_earned = 0
        
        # 保存更改
        db.session.commit()
        
        return jsonify({
            'message': '已成功标记为异常提交',
            'points_deducted': points_to_deduct,
            'submission': {
                'id': submission.id,
                'is_abnormal': True,
                'abnormal_reason': reason,
                'marked_at': submission.marked_at.isoformat(),
                'points_earned': 0
            },
            'participant': {
                'id': participant.id,
                'name': participant.name,
                'points_earned': participant.points_earned,
                'total_points_for_task': participant.total_points_for_task
            }
        }), 200
        
    except Exception as e:
        # 发生错误时回滚
        db.session.rollback()
        return jsonify({'error': f'标记异常提交失败: {str(e)}'}), 500

# 添加静态文件服务路由
@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    """提供上传文件的访问"""
    return send_from_directory('uploads', filename)

# 上传鼓励图片
@app.route('/api/admin/upload-encouragement-image', methods=['POST'])
@station_admin_required
def upload_encouragement_image(current_user_id):
    """上传鼓励内容图片"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': '没有上传图片文件'}), 400
            
        file = request.files['image']
        
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
            
        if file and allowed_file(file.filename):
            # 创建上传目录
            upload_folder = os.path.join('uploads', 'encouragement')
            os.makedirs(upload_folder, exist_ok=True)
            
            # 生成安全的文件名
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            
            # 格式：日期_时间_随机ID_原文件名
            new_filename = f"{timestamp}_{unique_id}_{filename}"
            file_path = os.path.join(upload_folder, new_filename)
            
            # 保存文件
            file.save(file_path)
            
            # 返回相对路径
            relative_path = os.path.join('uploads', 'encouragement', new_filename).replace('\\', '/')
            
            return jsonify({
                'success': True,
                'image_url': f'/{relative_path}'
            })
        else:
            return jsonify({'error': '不支持的文件类型'}), 400
            
    except Exception as e:
        print(f"上传鼓励图片失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'上传鼓励图片失败: {str(e)}'}), 500

# 更新全局鼓励设置
@app.route('/api/admin/global-encouragement-settings', methods=['PUT'])
@station_admin_required
def update_global_encouragement_settings(current_user_id):
    """更新全局鼓励设置"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': '数据格式不正确'}), 400
            
        # 获取全局设置
        global_settings = GlobalSettings.query.first()
        
        if not global_settings:
            # 如果不存在则创建
            global_settings = GlobalSettings()
            db.session.add(global_settings)
            
        # 更新鼓励信息
        global_settings.default_encouragement_message = data.get('default_encouragement_message', '')
        global_settings.default_encouragement_image_url = data.get('default_encouragement_image_url')
        
        # 提交到数据库
        db.session.commit()
        
        return jsonify({
            'success': True,
            'settings': global_settings.to_dict()
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"更新全局鼓励设置失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'更新全局鼓励设置失败: {str(e)}'}), 500

@app.route('/api/admin/daily-report/generate-word', methods=['POST'])
@station_admin_required
def generate_word_report(current_user_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Missing JSON in request"}), 400

        leaderboard_title = data.get('leaderboardTitle')
        leaderboard_data = data.get('leaderboardData') # 预期是 [{rank, nickname, score}, ...]
        invite_code_filter = data.get('inviteCode') # 邀请码筛选
        query_context = data.get('queryContext') # { type: "daily"|"custom_range"|"single_task", details: { startDate, endDate } | { taskName } }

        if not leaderboard_title or leaderboard_data is None: # leaderboard_data可以是空列表
            return jsonify({"error": "Missing leaderboardTitle or leaderboardData"}), 400
        if not isinstance(leaderboard_data, list):
            return jsonify({"error": "leaderboardData must be a list"}), 400

        # 导入word相关库
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import io

        document = Document()

        # Set default font to SimSun (宋体)
        style = document.styles['Normal']
        font = style.font
        font.name = 'SimSun' 
        # Ensure the font is applied to East Asian characters explicitly
        r = style.element.rPr
        # Create rFonts element if it doesn't exist
        if r.rFonts is None:
            r.get_or_add_rFonts()
        r.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r.rFonts.set(qn('w:ascii'), 'SimSun') # Also set for ASCII to ensure consistency if needed
        r.rFonts.set(qn('w:hAnsi'), 'SimSun') # And for HAnsi

        # 1. 页眉/页脚 (简单实现 - 生成时间)
        document.add_paragraph(f"日报生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph()

        # 2. 固定开头感谢语
        document.add_paragraph("感谢大家在本次活动中的热情参与和辛勤付出！")
        document.add_paragraph()

        # 3. 查询条件回顾
        query_recap = "本期日报统计条件："
        if invite_code_filter:
            query_recap += f"邀请码 '{invite_code_filter}'"
        
        if query_context:
            context_type = query_context.get('type')
            details = query_context.get('details', {})
            if invite_code_filter and context_type: query_recap += "，"

            if context_type == 'daily':
                query_recap += f"日期 {details.get('date', '今日')}"
            elif context_type == 'custom_range':
                query_recap += f"日期范围 {details.get('startDate', '?')} 至 {details.get('endDate', '?')}"
            elif context_type == 'single_task_active' or context_type == 'single_task_completed':
                status = "(进行中)" if context_type == 'single_task_active' else "(已结算)"
                query_recap += f"单活动 '{details.get('taskName', '未知活动')}' {status}"
            else:
                 if invite_code_filter: # 如果只有邀请码，没有其他上下文
                    pass # 邀请码信息已添加
                 elif not invite_code_filter and not context_type: # 如果什么上下文都没有
                    query_recap = "本期日报为通用统计。" # 提供一个通用回溯
                 else: # 其他未知类型
                    query_recap += "特定查询"

        else: # 兼容旧版可能只传了 invite_code 的情况，或者完全没有上下文
            if not invite_code_filter:
                query_recap = "本期日报为通用统计。"
        
        document.add_paragraph(query_recap)
        document.add_paragraph()

        # 4. 排行榜标题
        document.add_heading(leaderboard_title, level=1)

        # 5. 排行榜数据 (修改为纯文本列表)
        if leaderboard_data:
            document.add_paragraph("\n--- 排行榜详情 ---") # 添加一个分隔提示
            for item in leaderboard_data:
                rank = item.get('rank', '?')
                nickname = item.get('nickname', '未知用户')
                score = item.get('score', 'N/A')
                # Format: 第 X 名: @昵称 - YYY 积分
                paragraph_text = f"第 {rank} 名: @{nickname} - {score} 积分"
                document.add_paragraph(paragraph_text)
            document.add_paragraph("--- 排行榜结束 ---\n") # 添加一个分隔提示
        else:
            document.add_paragraph("本次统计暂无排行榜数据。")
        
        # document.add_paragraph() # 原有的空行，根据新格式可能不再需要，或调整

        # 6. 参与人数总览
        participant_count = len(leaderboard_data)
        document.add_paragraph(f"本次共有 {participant_count} 位粉丝上榜。")
        document.add_paragraph()

        # 7. 固定结尾鼓励语
        document.add_paragraph("再次恭喜榜上有名的各位！希望大家再接再厉，未来活动中争取更好成绩！也感谢所有参与的粉丝，你们的支持是我们最大的动力！")

        # 保存到内存流
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"daily_report_{timestamp}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error generating Word report: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "生成日报失败", "details": str(e)}), 500

@app.route('/api/station/tasks/<string:task_id>/settle', methods=['POST'])
@station_admin_required
def station_settle_task(current_user_id, task_id):
    """结算任务，将任务状态设为已完成"""
    # 查询任务并验证权限
    task = Task.query.join(Station).filter(
        Task.id == task_id,
        Station.owner_id == current_user_id
    ).first()
    
    if not task:
        return jsonify({"error": "任务不存在或您没有权限操作"}), 404
    
    # 检查任务是否已经是完成状态
    if task.status == 'completed':
        # 获取参与人数
        participant_count = Participant.query.filter_by(task_id=task.id).count()
        task_dict = task.to_dict()
        task_dict['participant_count'] = participant_count
        return jsonify({"message": "任务已经是结算状态", "task": task_dict}), 200
    
    try:
        # 更新任务状态为已完成
        task.status = 'completed'
        task.completed_at = datetime.utcnow()
        db.session.commit()
        
        # 获取参与人数
        participant_count = Participant.query.filter_by(task_id=task.id).count()
        task_dict = task.to_dict()
        task_dict['participant_count'] = participant_count
        
        return jsonify({
            "message": "任务结算成功",
            "task": task_dict
        }), 200
    except Exception as e:
        db.session.rollback()
        print(f"结算任务失败: {str(e)}")
        return jsonify({"error": f"结算任务失败: {str(e)}"}), 500

@app.route('/api/station/rankings', methods=['GET'])
@station_admin_required
def station_get_rankings(current_user_id):
    """根据查询参数获取排行榜数据"""
    # 获取查询参数
    ranking_type = request.args.get('type', 'overall')
    invite_code = request.args.get('invite_code')
    start_date = request.args.get('startDate')
    end_date = request.args.get('endDate')
    
    if not invite_code:
        return jsonify({
            "success": False,
            "message": "缺少必要参数：邀请码"
        }), 400
    
    try:
        # 查找邀请码
        invite_code_obj = InviteCode.query.filter_by(code=invite_code).first()
        if not invite_code_obj:
            return jsonify({
                "success": False,
                "message": "无效的邀请码"
            }), 404
        
        # 验证权限(检查邀请码是否属于管理员的站点)
        admin_stations = Station.query.filter_by(owner_id=current_user_id).all()
        admin_station_ids = [station.id for station in admin_stations]
        
        if invite_code_obj.station_id not in admin_station_ids:
            return jsonify({
                "success": False,
                "message": "无权访问此邀请码的排行榜"
            }), 403
        
        # 获取邀请码对应的站点ID
        station_id = invite_code_obj.station_id
        
        # 根据类型处理不同的排行榜数据
        if ranking_type == 'custom_range' and start_date and end_date:
            # 自定义日期范围排行榜
            try:
                start_datetime = datetime.strptime(start_date, '%Y-%m-%d')
                end_datetime = datetime.strptime(end_date, '%Y-%m-%d') + timedelta(days=1) # 包含结束日期
                
                # 获取日期范围内的参与数据
                rankings = db.session.query(
                    Participant.name.label('nickname'),
                    func.sum(Participant.points_earned).label('score'),
                    func.count(distinct(Participant.id)).label('completed_tasks')
                ).join(
                    Submission, Submission.participant_id == Participant.id
                ).join(
                    Task, Task.id == Participant.task_id
                ).filter(
                    Task.station_id == station_id,
                    Submission.submitted_at >= start_datetime,
                    Submission.submitted_at < end_datetime
                ).group_by(
                    Participant.name
                ).order_by(
                    func.sum(Participant.points_earned).desc()
                ).all()
                
                # 构建排行榜数据
                leaderboard_data = []
                for rank, (nickname, score, completed_tasks) in enumerate(rankings, 1):
                    leaderboard_data.append({
                        'rank': rank,
                        'nickname': nickname,
                        'score': score,
                        'completed_tasks': completed_tasks
                    })
                
                # 构建响应
                return jsonify({
                    "success": True,
                    "data": leaderboard_data,
                    "leaderboard_title": f"{start_date} 至 {end_date} 排行榜"
                })
                
            except ValueError as e:
                return jsonify({
                    "success": False,
                    "message": f"日期格式无效: {str(e)}"
                }), 400
        else:
            # 其他类型的排行榜...暂不实现
            return jsonify({
                "success": False,
                "message": "不支持的排行榜类型或缺少必要参数"
            }), 400
                
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"获取排行榜失败: {str(e)}"
        }), 500

@app.route('/api/station/tasks/<string:task_id>/ranking', methods=['GET'])
@station_admin_required
def station_get_task_ranking(current_user_id, task_id):
    """获取指定任务的排行榜"""
    # 获取任务
    task = Task.query.filter_by(id=task_id).first()
    if not task:
        return jsonify({
            "success": False,
            "message": "任务不存在"
        }), 404
        
    # 验证权限(检查任务是否属于管理员的站点)
    admin_stations = Station.query.filter_by(owner_id=current_user_id).all()
    admin_station_ids = [station.id for station in admin_stations]
    
    if task.station_id not in admin_station_ids:
        return jsonify({
            "success": False,
            "message": "无权访问此任务的排行榜"
        }), 403
    
    try:
        # 获取任务的参与者排行榜
        rankings = db.session.query(
            Participant.name.label('nickname'),
            Participant.points_earned.label('score'),
            Participant.submission_count.label('completed_tasks')
        ).filter(
            Participant.task_id == task_id
        ).order_by(
            Participant.points_earned.desc()
        ).all()
        
        # 构建排行榜数据
        leaderboard_data = []
        for rank, (nickname, score, completed_tasks) in enumerate(rankings, 1):
            leaderboard_data.append({
                'rank': rank,
                'nickname': nickname,
                'score': score,
                'completed_tasks': completed_tasks
            })
        
        # 构建响应
        return jsonify({
            "success": True,
            "data": leaderboard_data,
            "leaderboard_title": f"{task.title} 活动排行榜"
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False, 
            "message": f"获取任务排行榜失败: {str(e)}"
        }), 500

@app.route('/api/station/daily-report/generate-word', methods=['POST'])
@station_admin_required
def station_generate_word_report(current_user_id):
    # 实现和上面的逻辑相同，只是针对站点管理员
    try:
        # 直接调用处理请求，无需额外参数
        return generate_word_report(current_user_id)
    except Exception as e:
        print(f"生成Word报告失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": "生成日报失败", 
            "details": str(e)
        }), 500

# 管理员端反馈提交接口
@app.route('/api/feedback', methods=['POST'])
@station_admin_required
def admin_submit_feedback(current_user_id):
    """站子管理员提交反馈"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': '未接收到数据'}), 400
        
        # 基本验证
        if 'feedback_type' not in data or 'content' not in data:
            return jsonify({'error': '反馈类型和内容为必填项'}), 400
        
        if not data['content'] or len(data['content'].strip()) < 5:
            return jsonify({'error': '反馈内容太短，请提供更详细的信息'}), 400
        
        # 获取当前用户
        user = User.query.get(current_user_id)
        if not user:
            return jsonify({'error': '无法识别当前用户'}), 401
        
        # 创建反馈记录
        feedback = Feedback(
            user_type=data.get('user_type', 'admin'),
            feedback_type=data['feedback_type'],
            content=data['content'],
            contact=data.get('contact'),
            user_id=current_user_id,
            nickname=user.username
        )
        
        db.session.add(feedback)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '反馈提交成功，感谢您的意见！'
        }), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"提交反馈异常: {str(e)}")
        return jsonify({'error': '提交反馈时发生错误，请稍后重试'}), 500

# 粉丝端反馈提交接口
@app.route('/api/fan/feedback', methods=['POST'])
def fan_submit_feedback():
    """粉丝提交反馈"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': '未接收到数据'}), 400
        
        # 基本验证
        if 'feedback_type' not in data or 'content' not in data:
            return jsonify({'error': '反馈类型和内容为必填项'}), 400
        
        if not data['content'] or len(data['content'].strip()) < 5:
            return jsonify({'error': '反馈内容太短，请提供更详细的信息'}), 400
        
        # 创建反馈记录
        feedback = Feedback(
            user_type='fan',
            feedback_type=data['feedback_type'],
            content=data['content'],
            contact=data.get('contact'),
            nickname=data.get('nickname'),
            invite_code=data.get('invite_code')
        )
        
        db.session.add(feedback)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '反馈提交成功，感谢您的意见！'
        }), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"粉丝提交反馈异常: {str(e)}")
        return jsonify({'error': '提交反馈时发生错误，请稍后重试'}), 500

# 管理员获取反馈列表接口
@app.route('/api/feedback', methods=['GET'])
@station_admin_required
def admin_get_feedback_list(current_user_id):
    """站子管理员获取反馈列表"""
    try:
        # 获取分页和过滤参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        feedback_type = request.args.get('type')
        user_type = request.args.get('user_type')
        status = request.args.get('status')
        
        # 构建查询
        query = Feedback.query
        
        # 应用过滤条件
        if feedback_type:
            query = query.filter(Feedback.feedback_type == feedback_type)
        if user_type:
            query = query.filter(Feedback.user_type == user_type)
        if status:
            query = query.filter(Feedback.status == status)
            
        # 按创建时间倒序排序
        query = query.order_by(Feedback.created_at.desc())
        
        # 执行分页查询
        paginated_feedback = query.paginate(page=page, per_page=per_page, error_out=False)
        
        # 格式化结果
        feedbacks = [feedback.to_dict() for feedback in paginated_feedback.items]
        
        return jsonify({
            'success': True,
            'total': paginated_feedback.total,
            'pages': paginated_feedback.pages,
            'current_page': page,
            'items': feedbacks
        }), 200
        
    except Exception as e:
        print(f"获取反馈列表异常: {str(e)}")
        return jsonify({'error': '获取反馈列表时发生错误，请稍后重试'}), 500

# 管理员回复反馈接口
@app.route('/api/feedback/<string:feedback_id>/reply', methods=['POST'])
@station_admin_required
def admin_reply_feedback(current_user_id, feedback_id):
    """站子管理员回复反馈"""
    try:
        data = request.get_json()
        if not data or 'reply' not in data:
            return jsonify({'error': '回复内容为必填项'}), 400
        
        # 获取反馈记录
        feedback = Feedback.query.get(feedback_id)
        if not feedback:
            return jsonify({'error': '未找到反馈记录'}), 404
        
        # 更新反馈状态和回复内容
        feedback.reply = data['reply']
        feedback.status = 'replied'
        feedback.reviewed_at = datetime.utcnow()
        feedback.reviewed_by = current_user_id
        
        db.session.commit()
        
        # 如果有联系方式，可以发送邮件通知用户(可选实现)
        
        return jsonify({
            'success': True,
            'message': '回复提交成功'
        }), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"回复反馈异常: {str(e)}")
        return jsonify({'error': '回复反馈时发生错误，请稍后重试'}), 500

# 获取单个反馈详情接口
@app.route('/api/feedback/<string:feedback_id>', methods=['GET'])
@station_admin_required
def admin_get_feedback_detail(current_user_id, feedback_id):
    """站子管理员获取单个反馈详情"""
    try:
        # 获取反馈记录
        feedback = Feedback.query.get(feedback_id)
        if not feedback:
            return jsonify({'error': '未找到反馈记录'}), 404
        
        # 将状态更新为已查看（如果之前是待处理状态）
        if feedback.status == 'pending':
            feedback.status = 'reviewed'
            feedback.reviewed_at = datetime.utcnow()
            feedback.reviewed_by = current_user_id
            db.session.commit()
        
        return jsonify({
            'success': True,
            'data': feedback.to_dict()
        }), 200
        
    except Exception as e:
        print(f"获取反馈详情异常: {str(e)}")
        return jsonify({'error': '获取反馈详情时发生错误，请稍后重试'}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True) 
